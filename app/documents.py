from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any
import json
import re

from app.config import OUTPUT_DIR


def safe_slug(value: str) -> str:
    value = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", str(value).strip())
    return value.strip("_")[:48] or "document"


def today_prefix() -> str:
    return datetime.now().strftime("%Y%m%d")


def save_markdown(content: str, document_type: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = OUTPUT_DIR / f"{timestamp}_{safe_slug(document_type)}.md"
    path.write_text(content, encoding="utf-8")
    return path


def save_json(data: dict[str, Any], document_type: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = OUTPUT_DIR / f"{timestamp}_{safe_slug(document_type)}.json"
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def package_to_markdown(package: dict[str, Any]) -> str:
    lines: list[str] = []
    meta = package.get("package_metadata", {})
    lines.append(f"# {meta.get('title', '单节课教学包')}")
    lines.append("")
    if package.get("evaluator_score") is not None:
        lines.append(f"**evaluator_score：** {package.get('evaluator_score')}/100")
        comments = package.get("evaluator_comments", [])
        if comments:
            lines.append("**evaluator_comments：** " + "；".join(map(str, comments)))
        lines.append("")
    append_value(lines, package)
    return "\n".join(lines)


def append_value(lines: list[str], value: Any, level: int = 2, label: str | None = None) -> None:
    if isinstance(value, dict):
        if label:
            lines.append(f"{'#' * min(level, 6)} {human_label(label)}")
            lines.append("")
        for key, val in value.items():
            append_value(lines, val, min(level + 1, 6), key)
    elif isinstance(value, list):
        if label:
            lines.append(f"{'#' * min(level, 6)} {human_label(label)}")
            lines.append("")
        for item in value:
            if isinstance(item, (dict, list)):
                append_value(lines, item, min(level + 1, 6))
            else:
                lines.append(f"- {item}")
        lines.append("")
    else:
        if label:
            lines.append(f"**{human_label(label)}：** {value}")
        else:
            lines.append(str(value))
        lines.append("")


def human_label(value: str) -> str:
    labels = {
        "teacher_version": "教师授课版",
        "student_handout": "学生讲义版",
        "homework_answer_key": "作业与答案版",
        "parent_feedback": "家长反馈版",
        "student_record_update": "学生档案更新",
        "lesson_positioning": "本节课定位",
        "source_basis": "资料依据说明",
    }
    return labels.get(value, value.replace("_", " "))


def get_docx_tools():
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，无法导出 Word。") from exc
    return Document, Pt, qn


def setup_doc(doc: Any) -> None:
    _, Pt, qn = get_docx_tools()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)


def add_doc_value(doc: Any, value: Any, level: int = 1, label: str | None = None) -> None:
    if isinstance(value, dict):
        if label:
            doc.add_heading(human_label(label), level=min(level, 3))
        for key, val in value.items():
            add_doc_value(doc, val, min(level + 1, 4), key)
    elif isinstance(value, list):
        if label:
            doc.add_heading(human_label(label), level=min(level, 3))
        for item in value:
            if isinstance(item, (dict, list)):
                add_doc_value(doc, item, min(level + 1, 4))
            else:
                doc.add_paragraph(str(item), style="List Bullet")
    else:
        if label:
            p = doc.add_paragraph()
            p.add_run(f"{human_label(label)}：").bold = True
            p.add_run(str(value))
        else:
            doc.add_paragraph(str(value))


def save_docx_from_value(value: Any, title: str, path: Path) -> Path:
    Document, _, _ = get_docx_tools()
    doc = Document()
    setup_doc(doc)
    doc.add_heading(title, level=1)
    add_doc_value(doc, value, level=2)
    doc.save(path)
    return path


def package_file_base(package: dict[str, Any]) -> str:
    meta = package.get("package_metadata", {})
    student = safe_slug(meta.get("student", "上海新八女生"))
    subjects = safe_slug("_".join(meta.get("subjects", [])) or "语文数学")
    return f"{today_prefix()}_{student}_{subjects}"


def save_lesson_package_documents(package: dict[str, Any]) -> dict[str, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    base = package_file_base(package)
    paths: dict[str, Path] = {}

    full_md = package_to_markdown(package)
    paths["full_markdown"] = save_markdown(full_md, "lesson_package_full")
    paths["full_json"] = save_json(package, "lesson_package_full")

    teacher_payload = {
        "本节课定位": package.get("lesson_positioning", {}),
        "教师授课版": package.get("teacher_version", {}),
        "作业与答案版": package.get("homework_answer_key", {}),
        "资料依据说明": package.get("source_basis", {}),
        "质量检查": {
            "evaluator_score": package.get("evaluator_score"),
            "evaluator_comments": package.get("evaluator_comments", []),
            "requires_human_review": package.get("requires_human_review"),
        },
    }
    student_payload = package.get("student_handout", {})
    parent_payload = {
        "家长反馈版": package.get("parent_feedback", {}),
        "下节课建议": package.get("student_record_update", {}).get("next_lesson_suggestions", []),
    }

    paths["teacher_docx"] = save_docx_from_value(
        teacher_payload,
        "教师授课版",
        OUTPUT_DIR / f"{base}_教师授课版.docx",
    )
    paths["student_docx"] = save_docx_from_value(
        student_payload,
        "学生讲义版",
        OUTPUT_DIR / f"{base}_学生讲义版.docx",
    )
    paths["parent_docx"] = save_docx_from_value(
        parent_payload,
        "家长反馈版",
        OUTPUT_DIR / f"{base}_家长反馈版.docx",
    )
    return paths

