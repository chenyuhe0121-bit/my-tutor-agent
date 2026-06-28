from __future__ import annotations

from pathlib import Path


def read_local_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    return ""


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[系统提示] 未安装 python-docx，无法读取 Word 文件。"
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[系统提示] 未安装 pypdf，无法读取 PDF 文件。"
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages[:20], start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[PDF 第 {index} 页]\n{text.strip()}")
    return "\n\n".join(pages)

